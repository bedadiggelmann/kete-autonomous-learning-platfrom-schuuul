import os
import openai
import streamlit as st
from dotenv import load_dotenv
from langchain.memory import ConversationBufferMemory
from langchain.memory.chat_message_histories import StreamlitChatMessageHistory
from langchain.chains import LLMChain
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.callbacks.base import BaseCallbackHandler
from docx import Document
import json

# Load environment variables
load_dotenv()

# Set OpenAI connection details
openai.api_key = os.getenv("OPENAI_API_KEY")

st.title("Arbeitsblatt Generator")


class StreamHandler(BaseCallbackHandler):
    def __init__(self, container):
        self.container = container
        self.text = ""

    def on_llm_new_token(self, token: str, **kwargs) -> None:
        self.text += token
        self.container.markdown(self.text)


def add_footer(doc, klasse, lehrperson, schule):
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_text = f"Schule: {schule} | Klasse: {klasse} | Lehrperson: {lehrperson}"
    footer_para.text = footer_text


def create_doc(filename, topic, questions, include_answers, klasse, lehrperson, schule):
    doc = Document()
    doc.add_heading(topic, level=1)
    for question in questions:
        doc.add_paragraph(question["question"])
        if include_answers:
            doc.add_paragraph(f"Lösung: {question['answer']}")
        else:
            doc.add_paragraph("Antwort: _______________________")
    add_footer(doc, klasse, lehrperson, schule)

    if not os.path.exists("documents"):
        os.makedirs("documents")
    file_path = os.path.join("documents", filename)
    doc.save(file_path)
    return os.path.relpath(file_path)


# Define a list of available question types
available_question_types = [
    "Wahr oder Falsch",
    "Kurzantwort",
    "Lückentext",
]
# Define a list of available subjects
subjects = ["Mathematik", "Natur, Mensch, Gesellschaft", "Medien und Informatik"]

# Define a Streamlit interface for input
school = st.text_input("Schule eingeben", value="")
class_name = st.text_input("Klasse eingeben", value="")
teacher = st.text_input("Lehrperson eingeben", value="")
grade = st.selectbox(
    "Klassenstufe auswählen", list(range(1, 7)), index=0, key="grade_select"
)
subject = st.selectbox("Fach auswählen", subjects, key="subject_select")
topic = st.text_input("Thema eingeben", value="Multiplikation")
num_questions = st.number_input("Anzahl der Fragen", min_value=0, value=5, step=1)
question_types = st.multiselect("Fragetypen auswählen", available_question_types)

# Define a template for generating the conversation with JSON structure in the output
template = """
History:
{history}
Human:
{human_input}
"""

# Initialize memory and prompt template
msgs = StreamlitChatMessageHistory()
memory = ConversationBufferMemory(memory_key="history", chat_memory=msgs)
prompt_template = PromptTemplate(
    input_variables=["history", "human_input"], template=template
)
llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0)
llm_chain = LLMChain(llm=llm, prompt=prompt_template, memory=memory)
stream_handler = StreamHandler(st.empty())

if st.button("Arbeitsblatt generieren"):
    human_input = (
        f"Bitte generiere ein Arbeitsblatt als JSON für das Fach {subject}, "
        f"Klassenstufe {grade}, zum Thema {topic} mit {num_questions} Fragen. "
        f"Nutze die folgenden Fragetypen: {', '.join(question_types)}. Die Struktur des Arbeitsblattes "
        f"sollte 'topic', 'grade', 'subject', 'num_questions', 'question_types' und 'questions' beinhalten, "
        f"wobei 'questions' ein Array von Fragen ist, jede mit 'type', 'question', 'options' (wenn anwendbar) und 'answer'."
    )

    # Run the LLM chain to generate the response using the constructed human_input string
    response = llm_chain.run(
        {
            "history": msgs.messages,
            "human_input": human_input,
        },
        callbacks=[stream_handler],
    )

    # Display the raw response in the GUI for debugging
    st.text("Raw LLM Response:")
    st.write(response)

    # Try to parse the response as JSON and generate the document
    try:
        worksheet_data = json.loads(response)

        # Generate documents with and answers and without answers and store the returned file paths
        try:
            answers_path = create_doc(
                "Arbeitsblatt_mit_Antworten.docx",
                worksheet_data["topic"],
                worksheet_data["questions"],
                True,
                class_name,
                teacher,
                school,
            )
            no_answers_path = create_doc(
                "Arbeitsblatt_ohne_Antworten.docx",
                worksheet_data["topic"],
                worksheet_data["questions"],
                False,
                class_name,
                teacher,
                school,
            )

            # Set the paths in session state
            st.session_state["path_with_answers"] = answers_path
            st.session_state["path_without_answers"] = no_answers_path

        except Exception as e:  # Catch all exceptions during document creation
            st.error(f"An error occurred while creating the documents: {e}")

    except json.JSONDecodeError as e:
        st.error(f"An error occurred while parsing the LLM response: {e}")


# Display download buttons if paths are available
if "path_with_answers" in st.session_state and st.session_state["path_with_answers"]:
    with open(st.session_state["path_with_answers"], "rb") as file_with_answers:
        st.download_button(
            label="Download Arbeitsblatt mit Antworten",
            data=file_with_answers,
            file_name=os.path.basename(st.session_state["path_with_answers"]),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

if (
    "path_without_answers" in st.session_state
    and st.session_state["path_without_answers"]
):
    with open(st.session_state["path_without_answers"], "rb") as file_without_answers:
        st.download_button(
            label="Download Arbeitsblatt ohne Antworten",
            data=file_without_answers,
            file_name=os.path.basename(st.session_state["path_without_answers"]),
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
